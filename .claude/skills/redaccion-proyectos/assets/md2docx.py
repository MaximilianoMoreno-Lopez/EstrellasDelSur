# -*- coding: utf-8 -*-
"""Convierte el markdown de una solicitud en un .docx estructurado.

Uso:
    python md2docx.py entrada.md salida.docx --titulo "EuroÁgora" \
        --subtitulo "Solicitud Erasmus+ KA154-YOU" --extra "Convocatoria 2026"

Reglas de conversión: "# " y "## " pasan a Heading 1 con salto de página antes
de cada "## " salvo el primero, "### " a Heading 2, "#### " a Heading 3, listas
con guion o número a listas de Word, **negrita** y *cursiva* dentro del párrafo.
"""
import argparse
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x0A, 0x16, 0x28)
TEAL = RGBColor(0x0D, 0x94, 0x88)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_runs(par, text):
    """Soporta **negrita** y *cursiva* simples."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*", text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            par.add_run(m.group(1)).bold = True
        else:
            par.add_run(m.group(2)).italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def build(md_path, out_path, title, subtitle, extra, pie):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    for lvl, size, color, bold in ((1, 17, NAVY, True), (2, 14, NAVY, True), (3, 11, TEAL, True)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = bold

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("\n" + title)
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = NAVY
    for line, size in ((subtitle, 14), (extra, 12), ("\n" + pie, 10)):
        if not line or not line.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.font.color.rgb = GRAY if size == 10 else NAVY
    doc.add_page_break()

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    first_h2 = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            if not first_h2:
                doc.add_page_break()
            first_h2 = False
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

    doc.save(out_path)
    print("OK", out_path)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("entrada")
    ap.add_argument("salida")
    ap.add_argument("--titulo", default="Solicitud")
    ap.add_argument("--subtitulo", default="")
    ap.add_argument("--extra", default="")
    ap.add_argument("--pie", default="Borrador para copiar y pegar en el formulario oficial")
    a = ap.parse_args()
    build(a.entrada, a.salida, a.titulo, a.subtitulo, a.extra, a.pie)
